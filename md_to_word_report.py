import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


TITLE_FONT = "\u65b9\u6b63\u5c0f\u6807\u5b8b\u7b80\u4f53"
HEADING_FONT = "\u9ed1\u4f53"
BODY_FONT = "\u4eff\u5b8b_GB2312"

SIZE_2_PT = 22
SIZE_SMALL_2_PT = 18
SIZE_3_PT = 16
LINE_SPACING_PT = 28
FIRST_LINE_INDENT_PT = 36


def resolve_type_scale(body_size):
    normalized = str(body_size).strip().lower()
    small_2_names = {"small2", "xiaoer", "xiao-2", "2-", "18"}
    size_3_names = {"3", "sanhao", "size3", "three", "16"}

    if normalized in small_2_names:
        return {
            "body": SIZE_SMALL_2_PT,
            "title": SIZE_2_PT,
            "heading_1": SIZE_SMALL_2_PT,
        }
    if normalized in size_3_names:
        return {
            "body": SIZE_3_PT,
            "title": SIZE_SMALL_2_PT,
            "heading_1": SIZE_SMALL_2_PT,
        }

    raise ValueError("body size must be small2/xiaoer/18 or 3/sanhao/16")


def set_east_asian_font(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def set_paragraph_format(paragraph, first_line_indent=True, centered=False):
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(LINE_SPACING_PT)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line_indent:
        fmt.first_line_indent = Pt(FIRST_LINE_INDENT_PT)
    else:
        fmt.first_line_indent = Pt(0)
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_normal_style(doc, body_size_pt):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    style.font.size = Pt(body_size_pt)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(LINE_SPACING_PT)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(FIRST_LINE_INDENT_PT)


def clean_inline_markdown(text):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.strip()


def add_text_paragraph(
    doc,
    text,
    font_name,
    size_pt,
    first_line_indent=True,
    centered=False,
    bold=False,
):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=first_line_indent, centered=centered)
    run = paragraph.add_run(text)
    set_east_asian_font(run, font_name, size_pt, bold=bold)
    return paragraph


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def convert_markdown(md_path, output_path=None, first_h1_as_title=True, body_size="small2"):
    md_path = Path(md_path).resolve()
    output_path = Path(output_path).resolve() if output_path else md_path.with_suffix(".docx")
    text = md_path.read_text(encoding="utf-8-sig")
    scale = resolve_type_scale(body_size)

    doc = Document()
    configure_page(doc)
    set_normal_style(doc, scale["body"])

    used_title = False
    fallback_title_needed = True

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            content = clean_inline_markdown(heading.group(2))

            if fallback_title_needed and (not first_h1_as_title or level != 1):
                add_text_paragraph(
                    doc,
                    md_path.stem,
                    TITLE_FONT,
                    scale["title"],
                    first_line_indent=False,
                    centered=True,
                )
                fallback_title_needed = False

            if level == 1 and first_h1_as_title and not used_title:
                add_text_paragraph(
                    doc,
                    content,
                    TITLE_FONT,
                    scale["title"],
                    first_line_indent=False,
                    centered=True,
                )
                used_title = True
                fallback_title_needed = False
            elif level in (1, 2):
                add_text_paragraph(doc, content, HEADING_FONT, scale["heading_1"])
            else:
                add_text_paragraph(doc, content, BODY_FONT, scale["body"], bold=True)
            continue

        if fallback_title_needed:
            add_text_paragraph(
                doc,
                md_path.stem,
                TITLE_FONT,
                scale["title"],
                first_line_indent=False,
                centered=True,
            )
            fallback_title_needed = False

        line = re.sub(r"^[-*+]\s+", "", line)
        line = clean_inline_markdown(line)
        if line:
            add_text_paragraph(doc, line, BODY_FONT, scale["body"])

    if fallback_title_needed:
        add_text_paragraph(
            doc,
            md_path.stem,
            TITLE_FONT,
            scale["title"],
            first_line_indent=False,
            centered=True,
        )

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(
        description="Convert a Markdown report to DOCX using the local brief style."
    )
    parser.add_argument("input", help="Markdown file path")
    parser.add_argument("-o", "--output", help="Output DOCX path. Defaults to the same name.")
    parser.add_argument(
        "--body-size",
        default="small2",
        choices=["small2", "xiaoer", "18", "3", "sanhao", "16"],
        help="Body font size: small2/xiaoer/18 or 3/sanhao/16. Default: small2.",
    )
    parser.add_argument(
        "--filename-title",
        action="store_true",
        help="Use the filename as the document title instead of treating the first # heading as title.",
    )
    args = parser.parse_args()

    output = convert_markdown(
        args.input,
        output_path=args.output,
        first_h1_as_title=not args.filename_title,
        body_size=args.body_size,
    )
    print(str(output))


if __name__ == "__main__":
    main()
