#!/usr/bin/env python3
"""Convert a strictly formatted Haidian weekly-report Markdown file to DOCX."""

from __future__ import annotations

import argparse
import os
import re
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable
from urllib.parse import unquote, urlparse

try:
    from PIL import Image
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError as exc:  # pragma: no cover - handled by the launcher
    raise SystemExit(
        "缺少转换依赖。请运行“一键转Word.cmd”，或先执行："
        "python -m pip install -r requirements.txt"
    ) from exc


TITLE_FONT = "\u65b9\u6b63\u5927\u6807\u5b8b\u7b80\u4f53"
HEADING_FONT = "\u9ed1\u4f53"
BODY_FONT = "\u4eff\u5b8b_GB2312"
TABLE_FONT = "Microsoft YaHei"
FOOTER_FONT = "Calibri"

TITLE_SIZE_PT = 22
HEADING_SIZE_PT = 16
BODY_SIZE_PT = 16
TABLE_BODY_SIZE_PT = 10
TABLE_HEADER_SIZE_PT = 11
FOOTER_SIZE_PT = 9
LINE_SPACING_PT = 28
FIRST_LINE_INDENT_PT = 32
TABLE_WRAP_SIDE_DISTANCE_TWIPS = 180
TABLE_VERTICAL_OFFSET_TWIPS = 344

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_VERTICAL_CM = 2.54
MARGIN_HORIZONTAL_CM = 3.18
FOOTER_DISTANCE_CM = 1.27
USABLE_WIDTH_CM = PAGE_WIDTH_CM - (2 * MARGIN_HORIZONTAL_CM)

STYLE_TITLE = "Haidian Report Title"
STYLE_BLANK = "Haidian Report Blank"
STYLE_HEADING = "Haidian Report Heading 1"
STYLE_BODY = "Haidian Report Body"
STYLE_TABLE_BODY = "Haidian Report Table Body"
STYLE_TABLE_HEADER = "Haidian Report Table Header"
STYLE_IMAGE = "Haidian Report Image"

H1_RE = re.compile(r"^#(?!#)\s+(.+?)\s*$")
DEEP_HEADING_RE = re.compile(r"^#{2,6}\s+")
LIKELY_SECTION_RE = re.compile(r"^[一二三四五六七八九十百]+、\S")
IMAGE_ONLY_RE = re.compile(r"^\s*!\[([^\]]*)\]\((.+)\)\s*$")
ANY_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
LINK_RE = re.compile(r"(?<!!)\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"(\*\*|__)(.+?)\1")
TABLE_SEPARATOR_CELL_RE = re.compile(r"^:?-{3,}:?$")


class ConversionError(RuntimeError):
    """A user-facing conversion or validation error."""


@dataclass
class Block:
    kind: str
    line_no: int
    text: str = ""
    alt: str = ""
    image_ref: str = ""
    headers: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


@dataclass
class ParsedMarkdown:
    title: str
    title_line: int
    blocks: list[Block]


def _format_error(line_no: int, message: str) -> ConversionError:
    return ConversionError(f"第 {line_no} 行：{message}")


def _is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    cells = re.split(r"(?<!\\)\|", stripped)
    return [cell.replace(r"\|", "|").strip() for cell in cells]


def _is_table_separator(line: str) -> bool:
    if not _is_table_row(line):
        return False
    cells = _split_table_row(line)
    return bool(cells) and all(TABLE_SEPARATOR_CELL_RE.fullmatch(c.replace(" ", "")) for c in cells)


def _join_paragraph_lines(lines: Iterable[str]) -> str:
    result = ""
    for raw in lines:
        part = raw.strip()
        if not part:
            continue
        if result and result[-1:].isascii() and result[-1:].isalnum() and part[:1].isascii() and part[:1].isalnum():
            result += " "
        result += part
    return result


def parse_markdown(md_path: Path) -> ParsedMarkdown:
    try:
        text = md_path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ConversionError("Markdown 必须保存为 UTF-8 编码。") from exc

    lines = text.splitlines()
    first_nonempty = next((i for i, line in enumerate(lines) if line.strip()), None)
    if first_nonempty is None:
        raise ConversionError("Markdown 内容为空。")

    title_match = H1_RE.fullmatch(lines[first_nonempty].strip())
    if not title_match:
        raise _format_error(first_nonempty + 1, "首个非空行必须是以“# ”开头的主标题。")

    title = title_match.group(1).strip()
    if not title:
        raise _format_error(first_nonempty + 1, "主标题不能为空。")

    blocks: list[Block] = []
    i = first_nonempty + 1
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        line_no = i + 1

        if not stripped:
            i += 1
            continue

        if DEEP_HEADING_RE.match(stripped):
            raise _format_error(line_no, "仅支持一级 Markdown 标题“# ”，不支持“##”及更深层级。")

        heading_match = H1_RE.fullmatch(stripped)
        if heading_match:
            heading_text = heading_match.group(1).strip()
            if not heading_text:
                raise _format_error(line_no, "一级标题不能为空。")
            blocks.append(Block(kind="heading", line_no=line_no, text=heading_text))
            i += 1
            continue

        if LIKELY_SECTION_RE.match(stripped):
            raise _format_error(line_no, "疑似章节标题但未添加“# ”。请改为“# 章节标题”。")

        image_match = IMAGE_ONLY_RE.fullmatch(stripped)
        if image_match:
            blocks.append(
                Block(
                    kind="image",
                    line_no=line_no,
                    alt=image_match.group(1).strip(),
                    image_ref=image_match.group(2).strip(),
                )
            )
            i += 1
            continue

        if i + 1 < len(lines) and _is_table_row(raw) and _is_table_separator(lines[i + 1]):
            headers = _split_table_row(raw)
            separator = _split_table_row(lines[i + 1])
            if len(headers) != len(separator):
                raise _format_error(line_no, "表头与分隔行的列数不一致。")
            rows: list[list[str]] = []
            i += 2
            while i < len(lines) and _is_table_row(lines[i]):
                row = _split_table_row(lines[i])
                if len(row) != len(headers):
                    raise _format_error(i + 1, f"表格应为 {len(headers)} 列，当前行为 {len(row)} 列。")
                rows.append(row)
                i += 1
            blocks.append(Block(kind="table", line_no=line_no, headers=headers, rows=rows))
            continue

        paragraph_lines = [raw]
        i += 1
        while i < len(lines):
            candidate = lines[i]
            candidate_stripped = candidate.strip()
            if not candidate_stripped:
                break
            if H1_RE.fullmatch(candidate_stripped) or DEEP_HEADING_RE.match(candidate_stripped):
                break
            if LIKELY_SECTION_RE.match(candidate_stripped):
                break
            if IMAGE_ONLY_RE.fullmatch(candidate_stripped):
                break
            if i + 1 < len(lines) and _is_table_row(candidate) and _is_table_separator(lines[i + 1]):
                break
            paragraph_lines.append(candidate)
            i += 1

        paragraph_text = _join_paragraph_lines(paragraph_lines)
        inline_image = ANY_IMAGE_RE.search(paragraph_text)
        if inline_image:
            raise _format_error(line_no, "图片必须单独占一行，不能与正文写在同一行。")
        blocks.append(Block(kind="paragraph", line_no=line_no, text=paragraph_text))

    parsed = ParsedMarkdown(title=title, title_line=first_nonempty + 1, blocks=blocks)
    _validate_images(parsed, md_path.parent)
    return parsed


def _resolve_image_path(reference: str, base_dir: Path) -> Path:
    ref = reference.strip()
    if ref.startswith("<") and ref.endswith(">"):
        ref = ref[1:-1].strip()
    ref = unquote(ref)
    parsed = urlparse(ref)

    if parsed.scheme.lower() in {"http", "https"}:
        raise ConversionError(f"不支持网络图片，请先下载到本地：{reference}")

    if parsed.scheme.lower() == "file":
        raw_path = parsed.path or parsed.netloc
        raw_path = raw_path.replace("/", os.sep)
        if os.name == "nt" and re.match(r"^\\[A-Za-z]:", raw_path):
            raw_path = raw_path[1:]
        path = Path(raw_path)
    else:
        path = Path(ref)
        if not path.is_absolute():
            path = base_dir / path

    return path.expanduser().resolve()


def _validate_images(parsed: ParsedMarkdown, base_dir: Path) -> None:
    for block in parsed.blocks:
        if block.kind != "image":
            continue
        try:
            path = _resolve_image_path(block.image_ref, base_dir)
        except ConversionError as exc:
            raise _format_error(block.line_no, str(exc)) from exc
        if not path.is_file():
            raise _format_error(block.line_no, f"图片不存在：{path}")


def _normalize_font_name(name: str) -> str:
    lowered = name.casefold()
    lowered = re.sub(r"\([^)]*\)", "", lowered)
    return re.sub(r"[\s_\-&]+", "", lowered)


def _installed_windows_fonts() -> list[str]:
    if os.name != "nt":
        return []
    try:
        import winreg
    except ImportError:
        return []

    names: list[str] = []
    keys = (
        (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
        (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
    )
    for root, subkey in keys:
        try:
            with winreg.OpenKey(root, subkey) as key:
                value_count = winreg.QueryInfoKey(key)[1]
                for index in range(value_count):
                    value_name, _, _ = winreg.EnumValue(key, index)
                    names.append(value_name)
        except OSError:
            continue
    return names


def validate_required_fonts() -> None:
    if os.name != "nt":
        return
    installed = [_normalize_font_name(name) for name in _installed_windows_fonts()]
    required = {
        TITLE_FONT: [TITLE_FONT],
        HEADING_FONT: [HEADING_FONT, "SimHei"],
        BODY_FONT: [BODY_FONT, "FangSong_GB2312"],
        "微软雅黑": [TABLE_FONT, "微软雅黑"],
        FOOTER_FONT: [FOOTER_FONT],
    }
    missing: list[str] = []
    for display_name, aliases in required.items():
        alias_norms = [_normalize_font_name(alias) for alias in aliases]
        if not any(any(alias in font for alias in alias_norms) for font in installed):
            missing.append(display_name)
    if missing:
        raise ConversionError("缺少排版所需字体：" + "、".join(missing))


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def _set_style_font(style, font_name: str, size_pt: float, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def _get_or_add_paragraph_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, BODY_FONT, BODY_SIZE_PT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    title = _get_or_add_paragraph_style(doc, STYLE_TITLE)
    _set_style_font(title, TITLE_FONT, TITLE_SIZE_PT)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.left_indent = Pt(0)
    title.paragraph_format.right_indent = Pt(0)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)

    blank = _get_or_add_paragraph_style(doc, STYLE_BLANK)
    _set_style_font(blank, BODY_FONT, BODY_SIZE_PT)
    blank.paragraph_format.first_line_indent = Pt(0)
    blank.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    blank.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    blank.paragraph_format.space_before = Pt(0)
    blank.paragraph_format.space_after = Pt(0)

    heading = _get_or_add_paragraph_style(doc, STYLE_HEADING)
    _set_style_font(heading, HEADING_FONT, HEADING_SIZE_PT)
    heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.first_line_indent = Pt(FIRST_LINE_INDENT_PT)
    heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    heading.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(0)
    heading.paragraph_format.keep_with_next = False
    heading.paragraph_format.keep_together = False

    body = _get_or_add_paragraph_style(doc, STYLE_BODY)
    _set_style_font(body, BODY_FONT, BODY_SIZE_PT)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Pt(FIRST_LINE_INDENT_PT)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    body.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.keep_with_next = False
    body.paragraph_format.keep_together = False

    table_body = _get_or_add_paragraph_style(doc, STYLE_TABLE_BODY)
    _set_style_font(table_body, TABLE_FONT, TABLE_BODY_SIZE_PT)
    table_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table_body.paragraph_format.first_line_indent = Pt(0)
    table_body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    table_body.paragraph_format.line_spacing = Pt(16)
    table_body.paragraph_format.space_before = Pt(0)
    table_body.paragraph_format.space_after = Pt(0)

    table_header = _get_or_add_paragraph_style(doc, STYLE_TABLE_HEADER)
    _set_style_font(table_header, TABLE_FONT, TABLE_HEADER_SIZE_PT, bold=True)
    table_header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_header.paragraph_format.first_line_indent = Pt(0)
    table_header.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    table_header.paragraph_format.line_spacing = Pt(18)
    table_header.paragraph_format.space_before = Pt(0)
    table_header.paragraph_format.space_after = Pt(0)

    image = _get_or_add_paragraph_style(doc, STYLE_IMAGE)
    _set_style_font(image, BODY_FONT, BODY_SIZE_PT)
    image.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image.paragraph_format.first_line_indent = Pt(0)
    image.paragraph_format.space_before = Pt(0)
    image.paragraph_format.space_after = Pt(0)

    # WPS/Word enables widow/orphan control by default.  The reference report
    # explicitly disables it, allowing a single final line on the next page.
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            style.paragraph_format.widow_control = False


def _configure_document_settings(doc: Document) -> None:
    settings = doc.settings._element

    character_spacing = settings.find(qn("w:characterSpacingControl"))
    if character_spacing is None:
        character_spacing = OxmlElement("w:characterSpacingControl")
        settings.append(character_spacing)
    character_spacing.set(qn("w:val"), "compressPunctuation")

    # Mark East Asian typography as Simplified Chinese so WPS applies the
    # punctuation-compression rule consistently.
    theme_font_lang = settings.find(qn("w:themeFontLang"))
    if theme_font_lang is None:
        theme_font_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_font_lang)
    theme_font_lang.set(qn("w:val"), "zh-CN")
    theme_font_lang.set(qn("w:eastAsia"), "zh-CN")


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(MARGIN_VERTICAL_CM)
    section.bottom_margin = Cm(MARGIN_VERTICAL_CM)
    section.left_margin = Cm(MARGIN_HORIZONTAL_CM)
    section.right_margin = Cm(MARGIN_HORIZONTAL_CM)
    section.footer_distance = Cm(FOOTER_DISTANCE_CM)


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _add_page_number(doc: Document) -> None:
    section = doc.sections[0]
    paragraph = section.footer.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = " PAGE "
    instruction_run._r.append(instruction)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run("1")

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)

    for run in (begin_run, instruction_run, separate_run, result_run, end_run):
        _set_run_font(run, FOOTER_FONT, FOOTER_SIZE_PT)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _plain_link_text(text: str) -> str:
    previous = None
    current = text
    while previous != current:
        previous = current
        current = LINK_RE.sub(lambda match: match.group(1), current)
    return current


def _unescape_inline(text: str) -> str:
    return re.sub(r"\\([\\`*_{}\[\]()#+\-.!|])", r"\1", text)


def _inline_spans(text: str) -> list[tuple[str, bool]]:
    text = _plain_link_text(text)
    spans: list[tuple[str, bool]] = []
    cursor = 0
    for match in BOLD_RE.finditer(text):
        if match.start() > cursor:
            spans.append((_unescape_inline(text[cursor : match.start()]), False))
        spans.append((_unescape_inline(match.group(2)), True))
        cursor = match.end()
    if cursor < len(text):
        spans.append((_unescape_inline(text[cursor:]), False))
    if not spans:
        spans.append((_unescape_inline(text), False))
    return [(value, bold) for value, bold in spans if value]


def _add_inline_runs(paragraph, text: str, font_name: str, size_pt: float, force_bold: bool = False) -> None:
    for value, is_bold in _inline_spans(text):
        run = paragraph.add_run(value)
        _set_run_font(run, font_name, size_pt, bold=(force_bold or is_bold))


def _add_title(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph(style=STYLE_TITLE)
    run = paragraph.add_run(_plain_link_text(title))
    _set_run_font(run, TITLE_FONT, TITLE_SIZE_PT)
    doc.add_paragraph(style=STYLE_BLANK)


def _add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style=STYLE_HEADING)
    _add_inline_runs(paragraph, text, HEADING_FONT, HEADING_SIZE_PT)


def _add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style=STYLE_BODY)
    _add_inline_runs(paragraph, text, BODY_FONT, BODY_SIZE_PT)


def _image_width_cm(path: Path) -> float:
    with Image.open(path) as image:
        width_px, height_px = image.size
    if width_px <= 0 or height_px <= 0:
        raise ConversionError(f"无法读取图片尺寸：{path}")
    aspect = width_px / height_px
    if aspect >= 1.4:
        return USABLE_WIDTH_CM
    if aspect <= 0.9:
        return min(5.5, USABLE_WIDTH_CM)
    natural_at_96_dpi = width_px / 96 * 2.54
    return min(max(natural_at_96_dpi, 4.0), USABLE_WIDTH_CM)


def _add_image(doc: Document, path: Path, alt: str) -> None:
    paragraph = doc.add_paragraph(style=STYLE_IMAGE)
    run = paragraph.add_run()
    try:
        inline_shape = run.add_picture(str(path), width=Cm(_image_width_cm(path)))
    except Exception as exc:
        raise ConversionError(f"无法插入图片：{path}\n{exc}") from exc
    description = alt or path.name
    inline_shape._inline.docPr.set("descr", description)
    inline_shape._inline.docPr.set("title", description)


def _cm_to_twips(value_cm: float) -> int:
    return round(value_cm / 2.54 * 1440)


def _table_column_widths(headers: list[str]) -> list[float]:
    normalized = [_plain_link_text(re.sub(r"\*\*|__", "", h)).strip() for h in headers]
    if len(headers) == 5 and normalized[0] == "序号" and normalized[-1] == "传播量":
        # Hidden precision keeps the total at 15.82 cm while Word rounds the
        # five displayed column widths to 1.40 / 2.81 / 8.00 / 2.10 / 1.52 cm.
        return [1.398, 2.808, 7.998, 2.098, 1.518]
    equal_width = USABLE_WIDTH_CM / len(headers)
    return [equal_width] * len(headers)


def _set_cell_margins(cell, top: int = 80, left: int = 100, bottom: int = 80, right: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "000000")


def _set_table_wrap_around(table) -> None:
    tbl_pr = table._tbl.tblPr
    floating = tbl_pr.first_child_found_in("w:tblpPr")
    if floating is None:
        floating = OxmlElement("w:tblpPr")
        tbl_pr.insert(0, floating)
    floating.set(qn("w:leftFromText"), str(TABLE_WRAP_SIDE_DISTANCE_TWIPS))
    floating.set(qn("w:rightFromText"), str(TABLE_WRAP_SIDE_DISTANCE_TWIPS))
    floating.set(qn("w:horzAnchor"), "page")
    floating.set(qn("w:tblpXSpec"), "center")
    floating.set(qn("w:vertAnchor"), "text")
    floating.set(qn("w:tblpY"), str(TABLE_VERTICAL_OFFSET_TWIPS))
    for obsolete_attr in ("tblpX", "tblpYSpec", "topFromText", "bottomFromText"):
        floating.attrib.pop(qn(f"w:{obsolete_attr}"), None)

    overlap = tbl_pr.first_child_found_in("w:tblOverlap")
    if overlap is None:
        overlap = OxmlElement("w:tblOverlap")
        floating.addnext(overlap)
    overlap.set(qn("w:val"), "never")


def _set_fixed_table_geometry(table, widths_cm: list[float]) -> None:
    widths_twips = [_cm_to_twips(width) for width in widths_cm]
    total_twips = sum(widths_twips)
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths_cm[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.get_or_add_tcW()
            tc_w.set(qn("w:w"), str(widths_twips[index]))
            tc_w.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def _fill_table_cell(cell, text: str, is_header: bool) -> None:
    paragraph = cell.paragraphs[0]
    _clear_paragraph(paragraph)
    paragraph.style = STYLE_TABLE_HEADER if is_header else STYLE_TABLE_BODY
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT
    _add_inline_runs(
        paragraph,
        text,
        TABLE_FONT,
        TABLE_HEADER_SIZE_PT if is_header else TABLE_BODY_SIZE_PT,
        force_bold=is_header,
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_margins(cell)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths_cm = _table_column_widths(headers)

    for column_index, header in enumerate(headers):
        _fill_table_cell(table.rows[0].cells[column_index], header, is_header=True)
    for row_index, row in enumerate(rows, start=1):
        for column_index, value in enumerate(row):
            _fill_table_cell(table.rows[row_index].cells[column_index], value, is_header=False)

    _set_repeat_table_header(table.rows[0])
    _set_fixed_table_geometry(table, widths_cm)
    _set_table_borders(table)
    _set_table_wrap_around(table)


def _disable_widow_control(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    for paragraph in paragraphs:
        paragraph.paragraph_format.widow_control = False


def build_document(parsed: ParsedMarkdown, md_path: Path) -> Document:
    doc = Document()
    _configure_page(doc)
    _configure_styles(doc)
    _configure_document_settings(doc)
    _add_page_number(doc)
    _add_title(doc, parsed.title)

    for block in parsed.blocks:
        if block.kind == "heading":
            _add_heading(doc, block.text)
        elif block.kind == "paragraph":
            _add_body_paragraph(doc, block.text)
        elif block.kind == "image":
            image_path = _resolve_image_path(block.image_ref, md_path.parent)
            _add_image(doc, image_path, block.alt)
        elif block.kind == "table":
            _add_table(doc, block.headers, block.rows)
        else:  # pragma: no cover - defensive guard
            raise ConversionError(f"未知内容块：{block.kind}")

    _disable_widow_control(doc)

    doc.core_properties.title = parsed.title
    doc.core_properties.subject = "海淀意识形态周报"
    return doc


def convert_markdown(md_path: Path, output_path: Path) -> Path:
    md_path = md_path.expanduser().resolve()
    output_path = output_path.expanduser().resolve()
    if not md_path.is_file():
        raise ConversionError(f"找不到 Markdown 文件：{md_path}")
    if md_path.suffix.casefold() != ".md":
        raise ConversionError("输入文件必须是 .md 格式。")

    validate_required_fonts()
    parsed = parse_markdown(md_path)
    doc = build_document(parsed, md_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temp_handle = tempfile.NamedTemporaryFile(
        prefix=f".{output_path.stem}.", suffix=".docx", dir=output_path.parent, delete=False
    )
    temp_path = Path(temp_handle.name)
    temp_handle.close()
    try:
        doc.save(temp_path)
        os.replace(temp_path, output_path)
    except Exception:
        temp_path.unlink(missing_ok=True)
        raise
    return output_path


def _choose_markdown_file() -> Path | None:
    import tkinter as tk
    from tkinter import filedialog

    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    selected = filedialog.askopenfilename(
        title="选择要转换的海淀周报 Markdown",
        filetypes=(("Markdown 文档", "*.md"), ("所有文件", "*.*")),
    )
    root.destroy()
    return Path(selected) if selected else None


def _confirm_overwrite(path: Path) -> bool:
    import tkinter as tk
    from tkinter import messagebox

    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    result = messagebox.askyesno("确认覆盖", f"文件已存在：\n{path}\n\n是否覆盖？", parent=root)
    root.destroy()
    return result


def _show_error(message: str) -> None:
    import tkinter as tk
    from tkinter import messagebox

    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    messagebox.showerror("转换失败", message, parent=root)
    root.destroy()


def _open_output(path: Path) -> None:
    if os.name == "nt":
        os.startfile(path)  # type: ignore[attr-defined]
    else:  # pragma: no cover - the toolkit targets Windows
        raise ConversionError(f"当前系统无法自动打开文件，请手动打开：{path}")


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="将海淀周报 Markdown 转换为指定格式的 Word 文档。")
    parser.add_argument("input", nargs="?", help="输入 Markdown 文件；省略时弹出文件选择窗口。")
    parser.add_argument("-o", "--output", help="输出 DOCX 路径；默认与 Markdown 同目录、同文件名。")
    parser.add_argument("--force", action="store_true", help="不询问并覆盖已存在的输出文件。")
    parser.add_argument("--no-open", action="store_true", help="转换成功后不自动打开 Word。")
    parser.add_argument("--interactive", action="store_true", help=argparse.SUPPRESS)
    return parser


def main(argv: list[str] | None = None) -> int:
    args = _build_parser().parse_args(argv)
    interactive = args.interactive or not args.input

    try:
        md_path = Path(args.input) if args.input else _choose_markdown_file()
        if md_path is None:
            return 0
        md_path = md_path.expanduser().resolve()
        output_path = Path(args.output).expanduser().resolve() if args.output else md_path.with_suffix(".docx")

        if output_path.exists() and not args.force:
            if interactive:
                if not _confirm_overwrite(output_path):
                    return 0
            else:
                raise ConversionError(f"输出文件已存在：{output_path}\n如需覆盖，请添加 --force。")

        result = convert_markdown(md_path, output_path)
        print(f"已生成：{result}")
        if not args.no_open:
            _open_output(result)
        return 0
    except ConversionError as exc:
        message = str(exc)
    except Exception as exc:  # pragma: no cover - last-resort UX guard
        message = f"发生未预期错误：{exc}"

    if interactive:
        try:
            _show_error(message)
        except Exception:
            print(message, file=sys.stderr)
    else:
        print(message, file=sys.stderr)
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
