#!/usr/bin/env python3
"""Convert a Markdown special report into the fixed DOCX template."""

from __future__ import annotations

import argparse
import os
import re
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError as exc:  # pragma: no cover - handled by launcher
    raise SystemExit("缺少转换依赖。请运行“一键转Word.cmd”，或先执行：python -m pip install -r requirements.txt") from exc


BODY_FONT = "仿宋_GB2312"
HEADING_FONT = "黑体"
FOOTER_FONT = "Calibri"
BODY_SIZE_PT = 16
LINE_SPACING_PT = 28
FIRST_LINE_INDENT_PT = 32
FOOTER_SIZE_PT = 9

PLACEHOLDER_TOC = "{{TOC}}"
PLACEHOLDER_BODY = "{{BODY}}"
DEFAULT_TEMPLATE = Path(__file__).resolve().parent / "templates" / "境外敏感信息专报模板.docx"

ENTRY_RE = re.compile(r"^##\s+(.+?)\s*$")
DEEP_HEADING_RE = re.compile(r"^#{3,6}\s+")
IMAGE_RE = re.compile(r"!\[[^\]]*\]\([^)]*\)")
TABLE_RE = re.compile(r"^\s*\|.*\|\s*$")
LINK_RE = re.compile(r"(?<!!)\[([^\]]+)\]\(([^)]+)\)")
BOLD_RE = re.compile(r"(\*\*|__)(.+?)\1")


class ConversionError(RuntimeError):
    """User-facing conversion or validation error."""


@dataclass
class Entry:
    line_no: int
    index: str
    media: str
    date: str
    title: str
    toc_page: str = ""
    paragraphs: list[str] = field(default_factory=list)


@dataclass
class ParsedReport:
    title: str = "境外敏感信息专报"
    issue: str = ""
    issue_date: str = ""
    toc_pages: str = "manual"
    entries: list[Entry] = field(default_factory=list)


def _format_error(line_no: int, message: str) -> ConversionError:
    return ConversionError(f"第 {line_no} 行：{message}")


def _join_paragraph_lines(lines: list[str]) -> str:
    result = ""
    for raw in lines:
        part = raw.strip()
        if not part:
            continue
        if result and result[-1:].isascii() and result[-1:].isalnum() and part[:1].isascii() and part[:1].isalnum():
            result += " "
        result += part
    return result


def _clean_inline(text: str) -> str:
    text = LINK_RE.sub(lambda match: match.group(1), text)
    text = BOLD_RE.sub(lambda match: match.group(2), text)
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\\([\\`*_{}\[\]()#+\-.!|])", r"\1", text)
    # Match the report convention: avoid accidental English/number spacing inside Chinese prose.
    text = re.sub(r"(?<=[一-鿿])\s+(?=[A-Za-z0-9])", "", text)
    text = re.sub(r"(?<=[A-Za-z0-9])\s+(?=[一-鿿])", "", text)
    return text.strip()


def _parse_front_matter(lines: list[str]) -> tuple[dict[str, str], int]:
    metadata: dict[str, str] = {}
    if not lines or lines[0].strip() != "---":
        return metadata, 0

    for index in range(1, len(lines)):
        stripped = lines[index].strip()
        if stripped == "---":
            return metadata, index + 1
        if not stripped or stripped.startswith("#"):
            continue
        if ":" not in stripped:
            raise _format_error(index + 1, "元数据必须使用 key: value 格式。")
        key, value = stripped.split(":", 1)
        metadata[key.strip()] = value.strip().strip('"').strip("'")
    raise _format_error(1, "元数据块缺少结束行 ---。")


def _parse_entry_heading(line: str, line_no: int) -> Entry:
    match = ENTRY_RE.fullmatch(line.strip())
    if not match:
        raise _format_error(line_no, "条目标题必须使用“## 序号｜媒体｜日期｜标题｜页码”格式。")
    parts = [part.strip() for part in re.split(r"[｜|]", match.group(1))]
    if len(parts) < 4:
        raise _format_error(line_no, "条目标题至少需要 4 项：序号、媒体、日期、标题。")
    if len(parts) > 5:
        raise _format_error(line_no, "条目标题最多 5 项：序号、媒体、日期、标题、目录页码。")
    if any(not value for value in parts[:4]):
        raise _format_error(line_no, "序号、媒体、日期、标题均不能为空。")
    toc_page = parts[4] if len(parts) == 5 else ""
    return Entry(line_no=line_no, index=parts[0], media=parts[1], date=parts[2], title=parts[3], toc_page=toc_page)


def parse_markdown(md_path: Path, toc_pages_override: str | None = None) -> ParsedReport:
    try:
        text = md_path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ConversionError("Markdown 必须保存为 UTF-8 编码。") from exc

    lines = text.splitlines()
    if not any(line.strip() for line in lines):
        raise ConversionError("Markdown 内容为空。")

    metadata, start = _parse_front_matter(lines)
    parsed = ParsedReport(
        title=metadata.get("title", ParsedReport.title),
        issue=metadata.get("issue", ""),
        issue_date=metadata.get("issue_date", metadata.get("date", "")),
        toc_pages=metadata.get("toc_pages", "manual").lower(),
    )
    if toc_pages_override:
        parsed.toc_pages = toc_pages_override
    if parsed.toc_pages not in {"manual", "blank", "estimate"}:
        raise ConversionError("toc_pages 只能是 manual、blank 或 estimate。")

    current: Entry | None = None
    paragraph_lines: list[str] = []
    paragraph_start = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if current is None or not paragraph_lines:
            paragraph_lines = []
            return
        paragraph = _clean_inline(_join_paragraph_lines(paragraph_lines))
        if paragraph:
            current.paragraphs.append(paragraph)
        paragraph_lines = []

    for offset, raw in enumerate(lines[start:], start=start + 1):
        stripped = raw.strip()
        if not stripped:
            flush_paragraph()
            continue
        if IMAGE_RE.search(stripped):
            raise _format_error(offset, "v1 暂不支持图片，请删除图片后转换。")
        if TABLE_RE.match(stripped):
            raise _format_error(offset, "v1 暂不支持表格，请改为正文段落。")
        if DEEP_HEADING_RE.match(stripped):
            raise _format_error(offset, "v1 仅支持“##”条目标题，不支持“###”及更深标题。")
        if stripped.startswith("# "):
            # Allow a document title heading before entries; it does not generate content.
            if current is not None:
                raise _format_error(offset, "正文中不应再出现一级标题。")
            parsed.title = _clean_inline(stripped[2:].strip()) or parsed.title
            continue
        if ENTRY_RE.fullmatch(stripped):
            flush_paragraph()
            if current is not None:
                if not current.paragraphs:
                    raise _format_error(current.line_no, "该条目没有正文段落。")
                parsed.entries.append(current)
            current = _parse_entry_heading(stripped, offset)
            continue
        if current is None:
            raise _format_error(offset, "正文必须写在“## 序号｜媒体｜日期｜标题｜页码”条目下。")
        if not paragraph_lines:
            paragraph_start = offset
        paragraph_lines.append(raw)

    flush_paragraph()
    if current is not None:
        if not current.paragraphs:
            raise _format_error(current.line_no, "该条目没有正文段落。")
        parsed.entries.append(current)
    if not parsed.entries:
        raise ConversionError("未找到任何“##”条目。")
    if parsed.toc_pages == "manual":
        for entry in parsed.entries:
            if not entry.toc_page:
                raise _format_error(entry.line_no, "manual 目录页码模式下，条目标题末尾必须填写页码。")
    return parsed


def _normalize_font_name(name: str) -> str:
    lowered = name.casefold()
    lowered = re.sub(r"\([^)]*\)", "", lowered)
    return re.sub(r"[\s_\-&]+", "", lowered)


def _installed_windows_fonts() -> list[str]:
    if os.name != "nt":
        return []
    try:
        import winreg
    except ImportError:
        return []
    names: list[str] = []
    for root, subkey in (
        (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
        (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts"),
    ):
        try:
            with winreg.OpenKey(root, subkey) as key:
                for index in range(winreg.QueryInfoKey(key)[1]):
                    value_name, _, _ = winreg.EnumValue(key, index)
                    names.append(value_name)
        except OSError:
            continue
    return names


def validate_required_fonts() -> None:
    if os.name != "nt":
        return
    installed = [_normalize_font_name(name) for name in _installed_windows_fonts()]
    required = {
        "仿宋_GB2312": [BODY_FONT, "FangSong_GB2312"],
        "黑体": [HEADING_FONT, "SimHei"],
    }
    missing: list[str] = []
    for display_name, aliases in required.items():
        alias_norms = [_normalize_font_name(alias) for alias in aliases]
        if not any(any(alias in font for alias in alias_norms) for font in installed):
            missing.append(display_name)
    if missing:
        raise ConversionError("缺少排版所需字体：" + "、".join(missing))


def _set_run_font(run, font_name: str, size_pt: float, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def _set_style_font(style, font_name: str, size_pt: float, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), font_name)


def _get_or_add_paragraph_style(doc: Document, name: str):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, BODY_FONT, BODY_SIZE_PT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    body = _get_or_add_paragraph_style(doc, "Special Report Body")
    _set_style_font(body, BODY_FONT, BODY_SIZE_PT)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Pt(FIRST_LINE_INDENT_PT)
    body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    body.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)

    meta = _get_or_add_paragraph_style(doc, "Special Report Meta")
    _set_style_font(meta, BODY_FONT, BODY_SIZE_PT)
    meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.first_line_indent = Pt(0)
    meta.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    meta.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(0)

    title = _get_or_add_paragraph_style(doc, "Special Report Article Title")
    _set_style_font(title, BODY_FONT, BODY_SIZE_PT, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.first_line_indent = Pt(0)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)

    toc = _get_or_add_paragraph_style(doc, "Special Report TOC")
    _set_style_font(toc, BODY_FONT, BODY_SIZE_PT)
    toc.paragraph_format.first_line_indent = Pt(0)
    toc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    toc.paragraph_format.line_spacing = Pt(LINE_SPACING_PT)
    toc.paragraph_format.space_before = Pt(0)
    toc.paragraph_format.space_after = Pt(0)


def _configure_document_settings(doc: Document) -> None:
    settings = doc.settings._element
    character_spacing = settings.find(qn("w:characterSpacingControl"))
    if character_spacing is None:
        character_spacing = OxmlElement("w:characterSpacingControl")
        settings.append(character_spacing)
    character_spacing.set(qn("w:val"), "compressPunctuation")
    theme_font_lang = settings.find(qn("w:themeFontLang"))
    if theme_font_lang is None:
        theme_font_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_font_lang)
    theme_font_lang.set(qn("w:val"), "zh-CN")
    theme_font_lang.set(qn("w:eastAsia"), "zh-CN")


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _fill_paragraph(paragraph, text: str, style_name: str, font_name: str = BODY_FONT, size_pt: int = BODY_SIZE_PT, bold: bool = False) -> None:
    _clear_paragraph(paragraph)
    paragraph.style = style_name
    run = paragraph.add_run(text)
    _set_run_font(run, font_name, size_pt, bold=bold)


def _insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p = new_p
    new_para._element = new_p
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def _find_placeholder(doc: Document, placeholder: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == placeholder:
            return paragraph
    raise ConversionError(f"模板缺少占位符：{placeholder}。请检查模板文件。")


def _toc_page(entry: Entry, mode: str, estimated: dict[str, str]) -> str:
    if mode == "manual":
        return entry.toc_page
    if mode == "blank":
        return "待核"
    return estimated.get(entry.index, "待核")


def _estimate_toc_pages(entries: list[Entry]) -> dict[str, str]:
    pages: dict[str, str] = {}
    current_page = 1
    chars_on_page = 0
    capacity = 900
    for entry in entries:
        pages[entry.index] = str(current_page)
        entry_chars = len(entry.media) + len(entry.date) + len(entry.title) + sum(len(p) for p in entry.paragraphs)
        chars_on_page += entry_chars
        while chars_on_page > capacity:
            current_page += 1
            chars_on_page -= capacity
    return pages


def _replace_toc(doc: Document, parsed: ParsedReport) -> None:
    placeholder = _find_placeholder(doc, PLACEHOLDER_TOC)
    estimated = _estimate_toc_pages(parsed.entries) if parsed.toc_pages == "estimate" else {}
    first_text = f"{parsed.entries[0].index}、{parsed.entries[0].title}……{_toc_page(parsed.entries[0], parsed.toc_pages, estimated)}"
    _fill_paragraph(placeholder, first_text, "Special Report TOC")
    cursor = placeholder
    for entry in parsed.entries[1:]:
        text = f"{entry.index}、{entry.title}……{_toc_page(entry, parsed.toc_pages, estimated)}"
        cursor = _insert_paragraph_after(cursor)
        _fill_paragraph(cursor, text, "Special Report TOC")


def _replace_body(doc: Document, parsed: ParsedReport) -> None:
    placeholder = _find_placeholder(doc, PLACEHOLDER_BODY)
    cursor = placeholder
    first = True
    for entry in parsed.entries:
        if not first:
            cursor = _insert_paragraph_after(cursor)
            _fill_paragraph(cursor, "", "Special Report Body")
            cursor = _insert_paragraph_after(cursor)
        first = False
        _fill_paragraph(cursor, f"{entry.media}：{entry.title}", "Special Report Article Title", bold=True)
        cursor = _insert_paragraph_after(cursor)
        _fill_paragraph(cursor, entry.date, "Special Report Meta")
        for paragraph_text in entry.paragraphs:
            cursor = _insert_paragraph_after(cursor)
            _fill_paragraph(cursor, paragraph_text, "Special Report Body")


def _set_section_page_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def _add_page_number(section) -> None:
    paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction = OxmlElement("w:instrText")
    instruction.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction.text = " PAGE \\* MERGEFORMAT "
    instruction_run._r.append(instruction)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run("1")

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)

    for run in (begin_run, instruction_run, separate_run, result_run, end_run):
        _set_run_font(run, FOOTER_FONT, FOOTER_SIZE_PT)


def build_document(parsed: ParsedReport, template_path: Path) -> Document:
    if not template_path.is_file():
        raise ConversionError(f"找不到模板文件：{template_path}")
    doc = Document(template_path)
    _configure_styles(doc)
    _configure_document_settings(doc)
    _replace_toc(doc, parsed)
    _replace_body(doc, parsed)
    if doc.sections:
        _set_section_page_start(doc.sections[-1], 1)
        _add_page_number(doc.sections[-1])
    doc.core_properties.title = parsed.title
    doc.core_properties.subject = "境外敏感信息专报"
    return doc


def convert_markdown(md_path: Path, output_path: Path, template_path: Path, toc_pages: str | None = None) -> Path:
    md_path = md_path.expanduser().resolve()
    output_path = output_path.expanduser().resolve()
    template_path = template_path.expanduser().resolve()
    if not md_path.is_file():
        raise ConversionError(f"找不到 Markdown 文件：{md_path}")
    if md_path.suffix.casefold() != ".md":
        raise ConversionError("输入文件必须是 .md 格式。")
    validate_required_fonts()
    parsed = parse_markdown(md_path, toc_pages_override=toc_pages)
    if parsed.toc_pages == "estimate":
        print("提示：目录页码为估算值，请打开 Word 后人工核对。")
    doc = build_document(parsed, template_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    temp_handle = tempfile.NamedTemporaryFile(prefix=f".{output_path.stem}.", suffix=".docx", dir=output_path.parent, delete=False)
    temp_path = Path(temp_handle.name)
    temp_handle.close()
    try:
        doc.save(temp_path)
        os.replace(temp_path, output_path)
    except Exception:
        temp_path.unlink(missing_ok=True)
        raise
    return output_path


def _choose_markdown_file() -> Path | None:
    import tkinter as tk
    from tkinter import filedialog
    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    selected = filedialog.askopenfilename(title="选择要转换的境外敏感信息专报 Markdown", filetypes=(("Markdown 文档", "*.md"), ("所有文件", "*.*")))
    root.destroy()
    return Path(selected) if selected else None


def _confirm_overwrite(path: Path) -> bool:
    import tkinter as tk
    from tkinter import messagebox
    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    result = messagebox.askyesno("确认覆盖", f"文件已存在：\n{path}\n\n是否覆盖？", parent=root)
    root.destroy()
    return result


def _show_error(message: str) -> None:
    import tkinter as tk
    from tkinter import messagebox
    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    messagebox.showerror("转换失败", message, parent=root)
    root.destroy()


def _open_output(path: Path) -> None:
    if os.name == "nt":
        os.startfile(path)  # type: ignore[attr-defined]
    else:  # pragma: no cover
        raise ConversionError(f"当前系统无法自动打开文件，请手动打开：{path}")


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="将境外敏感信息专报 Markdown 转换为固定模板 Word 文档。")
    parser.add_argument("input", nargs="?", help="输入 Markdown 文件；省略时弹出文件选择窗口。")
    parser.add_argument("-o", "--output", help="输出 DOCX 路径；默认与 Markdown 同目录、同文件名。")
    parser.add_argument("--template", default=str(DEFAULT_TEMPLATE), help="模板 DOCX 路径。")
    parser.add_argument("--toc-pages", choices=["manual", "blank", "estimate"], help="覆盖 Markdown 中的目录页码模式。")
    parser.add_argument("--force", action="store_true", help="不询问并覆盖已存在的输出文件。")
    parser.add_argument("--no-open", action="store_true", help="转换成功后不自动打开 Word。")
    parser.add_argument("--interactive", action="store_true", help=argparse.SUPPRESS)
    return parser


def main(argv: list[str] | None = None) -> int:
    args = _build_parser().parse_args(argv)
    interactive = args.interactive or not args.input
    try:
        md_path = Path(args.input) if args.input else _choose_markdown_file()
        if md_path is None:
            return 0
        md_path = md_path.expanduser().resolve()
        output_path = Path(args.output).expanduser().resolve() if args.output else md_path.with_suffix(".docx")
        if output_path.exists() and not args.force:
            if interactive:
                if not _confirm_overwrite(output_path):
                    return 0
            else:
                raise ConversionError(f"输出文件已存在：{output_path}\n如需覆盖，请添加 --force。")
        result = convert_markdown(md_path, output_path, Path(args.template), toc_pages=args.toc_pages)
        print(f"已生成：{result}")
        if not args.no_open:
            _open_output(result)
        return 0
    except ConversionError as exc:
        message = str(exc)
    except Exception as exc:  # pragma: no cover - last-resort UX guard
        message = f"发生未预期错误：{exc}"
    if interactive:
        try:
            _show_error(message)
        except Exception:
            print(message, file=sys.stderr)
    else:
        print(message, file=sys.stderr)
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
